
from turtle import width
from docx import Document
from docx.shared import Inches
import pyttsx3
import speech_recognition as sr 


#creating The speaking function 
def speak(text):
    pyttsx3.speak(text)

#creating the writting from user audio
def writting(question):
    speak(question)
    r= sr.Recognizer()
    with sr.Microphone() as source:
        audio= r.listen(source)
        try:
            text= r.recognize_google(audio)
            return text
        except:
            speak("Sorry i didnt get you try typing")
            text= typing(question)
            return text

#creating typing function when speech is not recognised

def typing(question):
    text= input(question)
    return text



document = Document()

########### Adding CV Picture ###########

document.add_picture(
    'user.png',width=Inches(2.0)
    )

########### Greeting User ########################

speak("Hello we are going to create your CV by asking you questions answer in speech if not recognised type in the terminal below... ")
speak("Lets get started.....")

########### Adding Details of the User###############

name= writting("What is your name ? ")
age= writting("What is your age ? ")
email= writting("What is your email ? ") 
phone= writting("What is your phone number ? ") 

document.add_paragraph(
   'Name: '+  name+ ' \n  Age: ' + age + ' years \n Email:  ' + email + ' \n Cell: '+ phone
 )

###############Adding About me to the Cv ##################

document.add_heading("About Me")
about= writting("Tell us about your self")
document.add_paragraph(about)


############### Adding work Experience ##############
document.add_heading("Experience")

p = document.add_paragraph()
company= writting("Name of Company you worked at? ")
from_date= writting("When did you start working? ")
to_date= writting("When did you leave working? ")
p.add_run('Name of Company: '+company+ ' \n').bold= True
p.add_run( 'From: '+ from_date + ' \t To: ' + to_date +'\n').italic= True

experience= writting("Describe your experience at " + company)
p.add_run('Description: '+experience+ '\n')


while True:
    more_experience = writting("Do you have more experience at organisations answer with Yes or No)")
    if more_experience.lower()=='yes':
        company= writting("Name of Company you worked at? ")
        from_date= writting("When did you start working? ")
        to_date= writting("When did you leave working? ")

        p.add_run('Name of Company: '+company+ ' \n').bold= True
        p.add_run( 'From: '+ from_date + ' \t To: ' + to_date +'\n').italic= True

        experience= writting("Describe your experience at " + company)

        p.add_run('Description: '+experience+ '\n')
    else:
        break


###############Adding Skills##################

document.add_heading("Skills")
skill= writting("Enter your skill" )
sp = document.add_paragraph(skill)
sp.style='List Bullet'

while True:
    more_skill = writting("Do you have more Skills (Yes or No)")
    if more_skill.lower()=='yes':
        skill= writting("Enter your skill" )
        sp = document.add_paragraph(skill)
        sp.style='List Bullet'
    else:
        break

#########Footer####################
section= document.sections[0]
footer= section.footer
p= footer.paragraphs[0]
p.text="Cv generate by Samuel Codex @2022"




#name of file
file= name+' cv.docx'

speak("Your document is complete if any corrections to be done correct from  the file named " + file)
speak("Thank you for using Samuel Codex cv boot")
#Saving document 


document.save(file)

